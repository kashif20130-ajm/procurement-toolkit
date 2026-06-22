from __future__ import annotations

from datetime import date
from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
WHITE = RGBColor(255, 255, 255)


def build_consulting_proposal_docx(client_profile: dict, assessments: dict) -> BytesIO:
    document = Document()
    _apply_document_setup(document)
    _apply_styles(document)
    _add_cover(document, client_profile, assessments)
    _add_executive_summary(document, client_profile, assessments)
    _add_consulting_proposal(document, assessments)
    _add_scope_of_work(document, assessments)
    _add_deliverables(document, assessments)
    _add_duration(document, assessments)
    _add_commercial_proposal(document, assessments)
    _add_assumptions(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def save_consulting_proposal_docx(path: str, client_profile: dict, assessments: dict) -> None:
    buffer = build_consulting_proposal_docx(client_profile, assessments)
    with open(path, "wb") as output:
        output.write(buffer.getvalue())


def _apply_document_setup(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "AI Company Assessment | Consulting Proposal"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Confidential"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = GRAY


def _apply_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.208


def _add_cover(document: Document, client_profile: dict, assessments: dict) -> None:
    client_name = client_profile.get("client_name") or "Client"
    consultant = client_profile.get("consultant_name") or "Consulting Team"
    assessment_date = client_profile.get("assessment_date") or date.today()

    document.add_paragraph()
    kicker = document.add_paragraph("CONSULTING PROPOSAL")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_run(kicker.runs[0], size=11, color=GRAY, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Assessment Findings to Implementation Roadmap")
    _format_run(run, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Prepared for {client_name}")
    _format_run(run, size=14, color=GRAY)

    document.add_paragraph()
    _add_metadata_table(
        document,
        [
            ("Prepared by", consultant),
            ("Date", str(assessment_date)),
            ("Assessment modules", ", ".join(assessments.keys()) or "Pending assessment completion"),
            ("Document purpose", "Proposal for remediation, improvement, and implementation support"),
        ],
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_executive_summary(document: Document, client_profile: dict, assessments: dict) -> None:
    document.add_heading("Executive Summary", level=1)
    client_name = client_profile.get("client_name") or "the client"
    summary = _proposal_summary(client_name, assessments)
    _add_callout(document, summary)

    priority_gaps = _top_gap_rows(assessments, limit=6)
    if not priority_gaps.empty:
        document.add_heading("Priority Improvement Themes", level=2)
        _add_table(
            document,
            priority_gaps[["Module", "Gap", "Priority", "Responsible Function", "Target Date"]],
            [1.25, 2.75, 0.9, 1.2, 0.9],
        )


def _add_consulting_proposal(document: Document, assessments: dict) -> None:
    document.add_heading("Consulting Proposal", level=1)
    document.add_paragraph(
        "We propose a structured improvement engagement to convert assessment findings into practical, auditable, "
        "and sustainable operating improvements. The engagement will prioritize high-risk gaps, establish clear "
        "ownership, strengthen evidence trails, and support management with implementation-ready tools and governance."
    )
    document.add_paragraph(
        "The proposed approach combines diagnostic validation, remediation planning, process design, evidence pack "
        "development, stakeholder workshops, and implementation support. The work will be sequenced to produce quick "
        "risk reduction while building repeatable capabilities for ongoing compliance and maturity improvement."
    )

    scores = _assessment_score_table(assessments)
    if not scores.empty:
        document.add_heading("Assessment Baseline", level=2)
        _add_table(document, scores, [2.8, 1.0, 1.2, 1.5])


def _add_scope_of_work(document: Document, assessments: dict) -> None:
    document.add_heading("Scope of Work", level=1)
    scope_rows = [
        ("Mobilization and validation", "Confirm priorities, stakeholders, data sources, and implementation constraints."),
        ("Gap remediation design", "Translate findings into detailed remediation workplans, controls, templates, and evidence requirements."),
        ("Process and governance enhancement", "Update policies, workflows, approval gates, accountability, reporting cadence, and escalation routes."),
        ("Evidence and reporting pack", "Create evidence registers, compliance trackers, management dashboards, and progress reporting templates."),
        ("Implementation support", "Facilitate workshops, coach responsible functions, and track closure of priority actions."),
    ]
    if assessments:
        scope_rows.append(("Module-specific remediation", f"Address findings across {', '.join(assessments.keys())}."))
    _add_table(document, pd.DataFrame(scope_rows, columns=["Workstream", "Scope"]), [2.0, 4.5])


def _add_deliverables(document: Document, assessments: dict) -> None:
    document.add_heading("Deliverables", level=1)
    deliverables = [
        ("Project charter and mobilization plan", "Confirmed scope, governance, timeline, roles, and working rhythm."),
        ("Validated findings register", "Consolidated assessment findings, evidence status, priority, and accountable owner."),
        ("Remediation roadmap", "Actionable workplan with risk, root cause, impact, recommendation, effort, target date, and owner."),
        ("Updated control and evidence templates", "Practical templates for policy, procedure, evidence capture, reporting, and governance."),
        ("Executive steering pack", "Progress dashboard, decisions required, risks, dependencies, and next-step recommendations."),
        ("Final implementation report", "Summary of completed actions, residual risk, and forward roadmap."),
    ]
    if any("procurement_radar" in data for data in assessments.values()):
        deliverables.append(("Procurement maturity radar", "Baseline and target maturity view across procurement capability dimensions."))
    _add_table(document, pd.DataFrame(deliverables, columns=["Deliverable", "Description"]), [2.2, 4.3])


def _add_duration(document: Document, assessments: dict) -> None:
    document.add_heading("Project Duration", level=1)
    duration = _estimate_duration(assessments)
    document.add_paragraph(
        f"The proposed engagement duration is {duration}. The timeline can be refined after mobilization based on "
        "stakeholder availability, document completeness, number of workshops, and the volume of remediation actions."
    )
    plan = pd.DataFrame(
        [
            ("1", "Mobilization and validation", "Kickoff, document review, stakeholder alignment, workplan confirmation"),
            ("2-3", "Remediation design", "Root-cause review, action planning, evidence requirements, templates"),
            ("4-6", "Implementation support", "Workshops, governance updates, process/control design, progress tracking"),
            ("7-8", "Reporting and handover", "Executive pack, final report, residual risk, next-phase roadmap"),
        ],
        columns=["Week", "Phase", "Key Activities"],
    )
    _add_table(document, plan, [0.8, 2.0, 3.7])


def _add_commercial_proposal(document: Document, assessments: dict) -> None:
    document.add_heading("Commercial Proposal", level=1)
    modules = max(1, len(assessments))
    complexity = _complexity_multiplier(assessments)
    base_fee = 45000
    fee = round(base_fee * modules * complexity / 500) * 500
    vat = round(fee * 0.05 / 100) * 100
    total = fee + vat

    commercial = pd.DataFrame(
        [
            ("Professional fees", f"AED {fee:,.0f}", "Fixed fee for agreed scope and deliverables."),
            ("VAT estimate", f"AED {vat:,.0f}", "Calculated at 5%; final tax treatment subject to invoice rules."),
            ("Estimated total", f"AED {total:,.0f}", "Excludes travel, translation, legal opinions, and third-party tools unless agreed."),
        ],
        columns=["Commercial Item", "Amount", "Notes"],
    )
    _add_table(document, commercial, [1.8, 1.2, 3.5])
    document.add_paragraph(
        "Commercials are indicative and should be confirmed following a scoping workshop. Fees assume timely access to "
        "stakeholders, relevant documents, and decision-makers."
    )


def _add_assumptions(document: Document) -> None:
    document.add_heading("Key Assumptions", level=1)
    assumptions = [
        "Client will nominate a single engagement sponsor and module-level process owners.",
        "Client will provide relevant documents, data extracts, and access to stakeholders in a timely manner.",
        "Legal, tax, or regulatory opinions are outside scope unless separately agreed.",
        "Commercials exclude travel, translation, third-party systems, and external assurance fees.",
    ]
    for item in assumptions:
        document.add_paragraph(item, style="List Bullet")


def _proposal_summary(client_name: str, assessments: dict) -> str:
    if not assessments:
        return (
            f"This proposal sets out an implementation engagement for {client_name} following completion of the AI "
            "company assessment. Once module findings are finalized, the scope, deliverables, duration, and commercial "
            "estimate can be calibrated to the confirmed gap profile."
        )

    modules = ", ".join(assessments.keys())
    high_priority = sum(
        len(data.get("gap_remediation_plan", pd.DataFrame()).query("Priority == 'High'"))
        for data in assessments.values()
        if isinstance(data.get("gap_remediation_plan"), pd.DataFrame)
    )
    return (
        f"The assessment findings for {client_name} indicate improvement opportunities across {modules}. "
        f"The proposed engagement focuses on closing priority gaps, strengthening governance and evidence, "
        f"and reducing implementation risk. The current gap profile includes {high_priority} high-priority "
        "remediation item(s), which should be addressed through a structured workplan with accountable owners, "
        "target dates, and executive oversight."
    )


def _top_gap_rows(assessments: dict, limit: int) -> pd.DataFrame:
    frames = []
    for module_name, data in assessments.items():
        plan = data.get("gap_remediation_plan")
        if isinstance(plan, pd.DataFrame) and not plan.empty:
            frame = plan.copy()
            frame.insert(0, "Module", module_name)
            frames.append(frame)
    if not frames:
        return pd.DataFrame()
    combined = pd.concat(frames, ignore_index=True)
    order = {"High": 0, "Medium": 1, "Low": 2}
    combined["_priority_order"] = combined["Priority"].map(order).fillna(9)
    return combined.sort_values(["_priority_order", "Target Date"]).drop(columns=["_priority_order"]).head(limit)


def _assessment_score_table(assessments: dict) -> pd.DataFrame:
    rows = []
    for module_name, data in assessments.items():
        rows.append(
            {
                "Module": module_name,
                "Score": f"{data.get('score', 0)} / 100",
                "Risk Level": data.get("risk_level", data.get("status", "Not assessed")),
                "AI Suggested Score": f"{data.get('ai_assessment_score', data.get('score', 0))} / 100",
            }
        )
    return pd.DataFrame(rows)


def _estimate_duration(assessments: dict) -> str:
    if not assessments:
        return "4 to 6 weeks"
    modules = len(assessments)
    high_priority = 0
    for data in assessments.values():
        plan = data.get("gap_remediation_plan")
        if isinstance(plan, pd.DataFrame) and not plan.empty:
            high_priority += int((plan["Priority"] == "High").sum())
    if modules >= 4 or high_priority >= 12:
        return "10 to 12 weeks"
    if modules >= 2 or high_priority >= 6:
        return "8 to 10 weeks"
    return "6 to 8 weeks"


def _complexity_multiplier(assessments: dict) -> float:
    high_priority = 0
    total_gaps = 0
    for data in assessments.values():
        plan = data.get("gap_remediation_plan")
        if isinstance(plan, pd.DataFrame) and not plan.empty:
            high_priority += int((plan["Priority"] == "High").sum())
            total_gaps += len(plan)
    if high_priority >= 12 or total_gaps >= 30:
        return 1.35
    if high_priority >= 6 or total_gaps >= 15:
        return 1.15
    return 1.0


def _add_metadata_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    _set_table_width(table, [1.7, 4.8])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        _shade_cell(cells[0], HEADER_FILL)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
        for run in cells[0].paragraphs[0].runs:
            _format_run(run, bold=True, color=NAVY)


def _add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    cell.text = text
    _shade_cell(cell, LIGHT_FILL)
    _set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def _add_table(document: Document, dataframe: pd.DataFrame, widths_in: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.autofit = False
    _set_table_width(table, widths_in)

    for column_index, column in enumerate(dataframe.columns):
        cell = table.rows[0].cells[column_index]
        cell.text = str(column)
        _shade_cell(cell, HEADER_FILL)
        _set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            _format_run(run, bold=True, color=NAVY, size=9.5)

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            cells[column_index].text = str(value)
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cells[column_index])
            for paragraph in cells[column_index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    _format_run(run, size=9)

    document.add_paragraph()


def _set_table_width(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths_in):
            row.cells[index].width = Inches(width)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _format_run(run, size: float = 11, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
