from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader


MAX_DOCUMENT_CHARS = 12000
MAX_TOTAL_CHARS = 45000


def extract_uploaded_file(uploaded_file) -> dict:
    file_bytes = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()
    extracted_text = extract_file_bytes(file_bytes, uploaded_file.name)

    return {
        "name": uploaded_file.name,
        "type": uploaded_file.type or suffix.replace(".", "").upper(),
        "size_kb": round(len(file_bytes) / 1024, 1),
        "characters_extracted": len(extracted_text),
        "content": _limit_text(extracted_text, MAX_DOCUMENT_CHARS),
    }


def extract_file_path(file_path: Path) -> str:
    return extract_file_bytes(file_path.read_bytes(), file_path.name)


def extract_file_bytes(file_bytes: bytes, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    if suffix == ".docx":
        return _extract_docx(file_bytes)
    if suffix in {".xlsx", ".xls"}:
        return _extract_excel(file_bytes)
    if suffix == ".doc":
        return (
            "Legacy .doc files are not directly readable by python-docx. "
            "Convert this file to .docx and upload it again for content extraction."
        )
    if suffix in {".txt", ".md"}:
        return file_bytes.decode("utf-8", errors="ignore")
    return "Unsupported file type."


def build_document_context(extracted_documents: list[dict]) -> str:
    sections = []
    remaining = MAX_TOTAL_CHARS

    for document in extracted_documents:
        if remaining <= 0:
            break

        heading = f"Document: {document['name']} ({document['type']}, {document['size_kb']} KB)"
        content = document.get("content", "")
        section = f"{heading}\n{content}"
        sections.append(section[:remaining])
        remaining -= len(sections[-1])

    if not sections:
        return "No uploaded document content is available."

    return "\n\n---\n\n".join(sections)


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    page_text = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_text.append(f"Page {page_number}\n{text}")
    return "\n\n".join(page_text).strip() or "No readable text found in PDF."


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            parts.append(f"Table {table_index}\n" + "\n".join(rows))

    return "\n\n".join(parts).strip() or "No readable text found in Word document."


def _extract_excel(file_bytes: bytes) -> str:
    workbook = pd.read_excel(BytesIO(file_bytes), sheet_name=None, dtype=str)
    parts = []

    for sheet_name, dataframe in workbook.items():
        dataframe = dataframe.fillna("")
        preview = dataframe.head(200)
        parts.append(f"Sheet: {sheet_name}\n{preview.to_string(index=False)}")

    return "\n\n".join(parts).strip() or "No readable tables found in Excel workbook."


def _limit_text(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[Content truncated for AI context window.]"
