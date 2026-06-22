from __future__ import annotations

import re
from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.kpi import KPI_TARGET_METADATA
from utils.kpi import format_kpi_maturity
from utils.kpi import format_kpi_target
from utils.kpi import kpi_risk_from_maturity
from utils.recommendation_engine import concise_finding_blocks


TRAFFIC_COLORS = {
    "dark_red": "8B0000",
    "red": "C00000",
    "orange": "F4B183",
    "amber": "FFC000",
    "yellow": "FFD966",
    "light_green": "A9D18E",
    "green": "00B050",
    "navy": "1F4E78",
    "blue": "5B9BD5",
    "grey": "D9E2F3",
    "light_grey": "F2F2F2",
    "white": "FFFFFF",
    "black": "000000",
}

PRIORITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}


def _add_table(document: Document, dataframe: pd.DataFrame, column_widths: list[float] | None = None) -> None:
    if dataframe.empty:
        document.add_paragraph("No rows available.")
        return

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_autofit(table)
    header_cells = table.rows[0].cells
    for index, column in enumerate(dataframe.columns):
        header_cells[index].text = str(column)
        _format_cell(header_cells[index], bold=True)
    _repeat_table_header(table.rows[0])

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            _format_cell(cells[index])

    if column_widths:
        for row in table.rows:
            for index, width in enumerate(column_widths[: len(row.cells)]):
                row.cells[index].width = Inches(width)


def _compact_evidence_matrix(evidence_matrix: pd.DataFrame) -> pd.DataFrame:
    compact_columns = [
        "Assessment Criteria",
        "Required Evidence",
        "Source Documents",
        "Evidence Strength",
        "Evidence Score 0-5",
        "Confidence Level",
        "Compliance Status",
        "Score Cap Applied",
        "Final Score Rationale",
    ]
    available_columns = [column for column in compact_columns if column in evidence_matrix.columns]
    return evidence_matrix[available_columns].copy()


def _compact_scoring_model(scoring_model: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in scoring_model.iterrows():
        rows.append(
            {
                "Assessment Criteria": row.get("Assessment Criteria", ""),
                "Weightage %": row.get("Weightage %", ""),
                "Score 0-5": row.get("Score 0-5", ""),
                "Risk Rating": row.get("Risk Rating", ""),
                "Evidence Summary": _limit_words(row.get("Score Rationale", ""), 30),
                "Score Cap": row.get("Score Cap Applied", ""),
            }
        )
    return pd.DataFrame(rows)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 11), ("Heading 3", 10), ("Heading 4", 9)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["navy"])


def _set_landscape(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _set_portrait(document: Document) -> None:
    current = document.sections[-1]
    if current.orientation == WD_ORIENT.PORTRAIT:
        return
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def _format_cell(cell, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.bold = bold


def _set_table_autofit(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.tblLayout
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_critical_issues_at_glance(document: Document, guidance: list[dict]) -> None:
    issues = _priority_guidance(guidance, limit=5)
    document.add_heading("Critical Issues at a Glance", level=3)
    if not issues:
        document.add_paragraph("No failed or partially compliant assessment criteria were identified.")
        return
    for item in issues:
        document.add_paragraph(str(item.get("Title", "Assessment criterion")), style="List Bullet")


def _add_executive_scorecard(document: Document, module_name: str, data: dict) -> None:
    document.add_heading("Executive Scorecard", level=3)
    scorecard = pd.DataFrame(
        [
            {"Metric": "Module", "Result": module_name},
            {"Metric": "Overall Score", "Result": f"{data.get('score', 0)} / 100"},
            {"Metric": "Risk Rating", "Result": data.get("risk_level", data.get("status", "Not assessed"))},
            {"Metric": "Assessment Mode", "Result": data.get("assessment_mode", "Not specified")},
        ]
    )
    _add_table(document, scorecard, column_widths=[2.0, 4.5])


def _add_risk_heat_map(document: Document, scoring_model: pd.DataFrame) -> None:
    document.add_heading("Risk Heat Map", level=3)
    counts = scoring_model.get("Risk Rating", pd.Series(dtype=str)).value_counts().to_dict()
    heat_map = pd.DataFrame(
        [
            {"Risk Rating": risk, "Criteria Count": counts.get(risk, 0)}
            for risk in ["Critical", "High", "Medium", "Low"]
        ]
    )
    _add_table(document, heat_map, column_widths=[2.0, 1.2])


def _add_compliance_dashboard(document: Document, evidence_matrix: pd.DataFrame) -> None:
    document.add_heading("Compliance Dashboard", level=3)
    counts = evidence_matrix.get("Compliance Status", pd.Series(dtype=str)).value_counts().to_dict()
    dashboard = pd.DataFrame(
        [
            {"Compliance Status": "Compliant", "Criteria Count": counts.get("Compliant", 0)},
            {"Compliance Status": "Partially Compliant", "Criteria Count": counts.get("Partially Compliant", 0)},
            {"Compliance Status": "Non-Compliant", "Criteria Count": counts.get("Non-Compliant", 0)},
        ]
    )
    _add_table(document, dashboard, column_widths=[2.2, 1.2])


def _add_concise_findings(document: Document, guidance: list[dict], scp_references: pd.DataFrame | None = None) -> None:
    scp_reference = _top_scp_reference(scp_references)
    for block in concise_finding_blocks(guidance):
        if block.get("Type") == "KPI Framework":
            continue
        document.add_heading(str(block["Title"]), level=4)
        _add_label_bullet(document, "Finding", block["Finding"])
        _add_label_bullet(document, "Risk", block["Risk"])
        recommendation = block["Recommendation"]
        if scp_reference:
            recommendation = _limit_words(f"Based on SCP reference: {scp_reference}, {recommendation}", 75)
        _add_label_bullet(document, "Recommendation", recommendation)
        _add_label_bullet(document, "Target Score", block["Target Score"])
        _add_label_bullet(document, "Priority", block["Priority"])
        _add_label_bullet(document, "Timeline", block["Timeline"])
        _add_label_bullet(document, "Required Documents", "; ".join(block.get("Required Documents", [])[:5]))


def _add_label_bullet(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(text))


def _add_limited_bullets(document: Document, items: list[str], limit: int) -> None:
    if not items:
        document.add_paragraph("No items available.")
        return
    for item in items[:limit]:
        document.add_paragraph(str(item), style="List Bullet")


def _evidence_summary_table(evidence_matrix: pd.DataFrame | None, evidence_register: pd.DataFrame | None) -> pd.DataFrame:
    if not isinstance(evidence_matrix, pd.DataFrame) or evidence_matrix.empty:
        return pd.DataFrame(columns=["Assessment Criteria", "Evidence Found", "Missing Evidence"])
    evidence_register = evidence_register if isinstance(evidence_register, pd.DataFrame) else pd.DataFrame()
    rows = []
    for _, row in evidence_matrix.iterrows():
        criterion = str(row.get("Assessment Criteria", ""))
        rows.append(
            {
                "Assessment Criteria": criterion,
                "Evidence Found": _evidence_found_summary(criterion, evidence_register, row),
                "Missing Evidence": _missing_evidence_summary(row),
            }
        )
    return pd.DataFrame(rows, columns=["Assessment Criteria", "Evidence Found", "Missing Evidence"])


def _evidence_found_summary(criterion: str, evidence_register: pd.DataFrame, row: pd.Series) -> str:
    labels = []
    if isinstance(evidence_register, pd.DataFrame) and not evidence_register.empty:
        criterion_lower = criterion.lower()
        for _, evidence in evidence_register.iterrows():
            affected = str(evidence.get("Affected Criteria", "")).lower()
            if criterion_lower in affected:
                labels.append(_short_evidence_label(str(evidence.get("Extract", ""))))
            if len(labels) >= 5:
                break
    if not labels:
        source_documents = str(row.get("Source Documents", "None"))
        if source_documents and source_documents != "None":
            labels = [item.strip() for item in source_documents.split(",")[:5] if item.strip()]
    labels = list(dict.fromkeys(labels))
    if not labels:
        return "No evidence found"
    return "\n".join(f"✓ {label}" for label in labels[:5])


def _missing_evidence_summary(row: pd.Series) -> str:
    text = str(row.get("Negative Evidence Found", "None"))
    required = str(row.get("Required Evidence", ""))
    if not text or text == "None":
        missing = _missing_required_labels(required, row)
        return "\n".join(f"✗ {label}" for label in missing[:5]) if missing else "None"
    labels = []
    for part in text.split("|"):
        label = part
        if "Negative evidence:" in label:
            label = label.split("Negative evidence:", 1)[1]
        if "Severe negative evidence:" in label:
            label = label.split("Severe negative evidence:", 1)[1]
        label = label.split("(")[0].strip(" .")
        if label and label not in labels:
            labels.append(label)
        if len(labels) >= 5:
            break
    return "\n".join(f"✗ {label}" for label in labels[:5]) if labels else "✗ Negative evidence identified"


def _short_evidence_label(extract: str) -> str:
    text = " ".join(str(extract).split())
    lowered = text.lower()
    mappings = [
        ("rfq", "RFQ"),
        ("quotation", "Quotation"),
        ("quote", "Quotation"),
        ("evaluation", "Evaluation Sheet"),
        ("purchase order", "Purchase Order"),
        ("goods received note", "GRN"),
        ("goods receipt", "GRN"),
        ("supplier assessment", "Supplier Assessment"),
        ("due diligence", "Supplier Due Diligence"),
        ("contract register", "Contract Register"),
        ("kpi", "KPI Record"),
        ("policy", "Policy"),
        ("procedure", "Procedure"),
        ("approval", "Approval Record"),
        ("reconciliation", "Reconciliation"),
        ("vat", "VAT Record"),
        ("emissions", "Emissions Record"),
        ("climate", "Climate Record"),
        ("esg", "ESG Record"),
    ]
    for needle, label in mappings:
        if needle in lowered:
            return label
    words = text.split()
    return " ".join(words[:8]) if words else "Evidence"


def _missing_required_labels(required_evidence: str, row: pd.Series) -> list[str]:
    score = int(row.get("Evidence Score 0-5", row.get("Score 0-5", 0)) or 0)
    if score >= 5:
        return []
    labels = []
    for part in str(required_evidence).replace(" and ", ", ").split(","):
        label = _short_evidence_label(part)
        if label and label not in labels:
            labels.append(label)
        if len(labels) == 5:
            break
    return labels


def _phase_roadmap(roadmap: pd.DataFrame) -> pd.DataFrame:
    if not isinstance(roadmap, pd.DataFrame) or roadmap.empty:
        return pd.DataFrame(columns=["Phase", "Timeline", "Priority Focus", "Actions"])
    phases = [
        ("Phase 1 (0-30 Days)", "1 Month"),
        ("Phase 2 (31-60 Days)", "3 Months"),
        ("Phase 3 (61-90 Days)", "6 Months"),
    ]
    rows = []
    for phase, timeline in phases:
        subset = roadmap[roadmap["Estimated Implementation Time"].astype(str).eq(timeline)]
        actions = subset["Area"].astype(str).head(8).tolist()
        priorities = ", ".join(sorted(set(subset["Priority"].astype(str)))) if not subset.empty else "None"
        rows.append(
            {
                "Phase": phase,
                "Timeline": _timeline_label(timeline),
                "Priority Focus": priorities,
                "Actions": "; ".join(actions) if actions else "No actions assigned.",
            }
        )
    remaining = roadmap[roadmap["Estimated Implementation Time"].astype(str).eq("12 Months")]
    if not remaining.empty:
        rows.append(
            {
                "Phase": "Phase 3 (61-90 Days)",
                "Timeline": _timeline_label("12 Months"),
                "Priority Focus": ", ".join(sorted(set(remaining["Priority"].astype(str)))),
                "Actions": "; ".join(remaining["Area"].astype(str).head(8).tolist()),
            }
        )
    return pd.DataFrame(rows)


def _compact_references(references: pd.DataFrame) -> pd.DataFrame:
    columns = [column for column in ["Module", "Source", "Chunk", "Relevance Score"] if column in references.columns]
    compact = references[columns].copy() if columns else references.copy()
    return compact.head(10)


def _top_scp_reference(references: pd.DataFrame | None) -> str:
    if not isinstance(references, pd.DataFrame) or references.empty:
        return ""
    if "Source" in references.columns:
        sources = [str(source) for source in references["Source"].dropna().head(2).tolist()]
        return "; ".join(dict.fromkeys(sources))
    return ""


def _compact_kpi_table(kpi_table: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in kpi_table.iterrows():
        status = str(row.get("Status", "Not Measured"))
        raw_score = row.get("KPI Score 0-5", pd.NA)
        score = None if pd.isna(raw_score) else int(raw_score)
        if status == "Not Measured" or score is None:
            priority = "Medium"
        else:
            priority = "High" if score <= 2 else "Medium" if score == 3 else "Low"
        rows.append(
            {
                "KPI": row.get("KPI Name", ""),
                "Current": row.get("Current", "Not Established"),
                "Reason": row.get(
                    "Current Value Reason",
                    "No KPI framework or KPI-specific evidence was identified in the uploaded documents.",
                ),
                "Status": status,
                "Target": _kpi_target_display(str(row.get("KPI Name", "")), row.get("Target", "")),
                "Owner": row.get("Owner", ""),
                "Priority": priority,
            }
        )
    return pd.DataFrame(rows, columns=["KPI", "Current", "Reason", "Status", "Target", "Owner", "Priority"])


def _kpi_formula_appendix(kpi_table: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in kpi_table.iterrows():
        rows.append(
            {
                "KPI": row.get("KPI Name", ""),
                "Formula": row.get("Formula", ""),
                "Data Source": row.get("Data Source", ""),
            }
        )
    return pd.DataFrame(rows, columns=["KPI", "Formula", "Data Source"])


def _maturity_radar_from_scoring_model(scoring_model: pd.DataFrame) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Dimension": row.get("Assessment Criteria", ""),
                "Maturity 1-5": row.get("Score 0-5", 0),
                "Risk Rating": row.get("Risk Rating", ""),
            }
            for _, row in scoring_model.iterrows()
        ]
    )


def _limit_words(text: object, max_words: int) -> str:
    words = str(text).split()
    if len(words) <= max_words:
        return " ".join(words)
    return " ".join(words[:max_words]).rstrip(" .,;:") + "."


def _sentence_key(text: object) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(text).lower()).strip()


def _split_report_sentences(text: object) -> list[str]:
    normalized = re.sub(r"\s+", " ", str(text or "")).strip()
    if not normalized or normalized.lower() in {"none", "nan", "not available."}:
        return []
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+|\s*;\s*", normalized) if part.strip()]


def _clean_report_text(text: object, max_sentences: int | None = None, max_words: int | None = None) -> str:
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    value = re.sub(
        r"Evidence for full compliance should demonstrate\s+(?:a\s+)?5/5 score would require",
        "Evidence should demonstrate",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        r"and should require management review before the\s*\.",
        "and should require documented management review.",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        r"current maturity is\s+(\d)/5\.\s*The current assessment score is\s+\1/5\.?",
        r"current score is \1/5.",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(r"\bbefore the\s*\.", "before approval.", value, flags=re.IGNORECASE)
    sentences = []
    seen = set()
    for sentence in _split_report_sentences(value):
        key = _sentence_key(sentence)
        if not key or key in seen:
            continue
        seen.add(key)
        sentences.append(sentence)
        if max_sentences and len(sentences) >= max_sentences:
            break
    result = " ".join(sentences)
    return _limit_words(result, max_words) if max_words else result


def _deduplicate_report_guidance(guidance: list[dict]) -> list[dict]:
    merged: dict[tuple[str, str], dict] = {}
    narrative_fields = [
        "Current State",
        "Gap Identified",
        "Risk Impact",
        "Required Actions",
        "Required Documents",
        "Required Controls",
        "Evidence Expected",
        "Target State",
        "Expected Score Improvement",
    ]
    for raw in guidance or []:
        if not isinstance(raw, dict):
            continue
        item = dict(raw)
        key = (str(item.get("Type", "Assessment Criterion")), _sentence_key(item.get("Title", "Assessment criterion")))
        for field in narrative_fields:
            item[field] = _clean_report_text(item.get(field, ""))
        existing = merged.get(key)
        if existing is None:
            merged[key] = item
            continue
        for field in narrative_fields:
            if len(str(item.get(field, ""))) > len(str(existing.get(field, ""))):
                existing[field] = item[field]
        for field in ["Current Score", "Target Score", "Priority", "Estimated Implementation Time", "Estimated Maturity Impact"]:
            if field in item and field not in existing:
                existing[field] = item[field]
    return list(merged.values())


def _priority_guidance(guidance: list[dict], limit: int = 5) -> list[dict]:
    assessment_items = [
        item
        for item in _deduplicate_report_guidance(guidance)
        if item.get("Type", "Assessment Criterion") == "Assessment Criterion"
        and int(item.get("Current Score", 0)) < int(item.get("Target Score", 5))
    ]
    return sorted(
        assessment_items,
        key=lambda item: (
            PRIORITY_ORDER.get(str(item.get("Priority", "Medium")), 2),
            int(item.get("Current Score", 0)),
            str(item.get("Title", "")),
        ),
    )[:limit]


def _business_summary(module_name: str, data: dict, guidance: list[dict]) -> str:
    score = int(data.get("score", 0) or 0)
    risk = str(data.get("risk_level", data.get("status", "Not assessed")))
    sentences = [f"The {module_name} achieved an overall score of {score}/100 with a {risk} risk rating."]
    raw_summary = data.get("executive_summary", data.get("analysis", ""))
    for sentence in _split_report_sentences(_clean_report_text(raw_summary, max_sentences=4)):
        lowered = sentence.lower()
        if any(
            term in lowered
            for term in [
                "overall score",
                "risk rating",
                "current maturity",
                "current assessment score",
                "demo mode",
                "without calling openai",
                "rule-based demo",
                "workflow testing",
                "api costs",
            ]
        ):
            continue
        sentences.append(_limit_words(sentence, 30))
        if len(sentences) == 4:
            break
    issues = _priority_guidance(guidance, limit=5)
    if issues:
        issue_names = ", ".join(str(item.get("Title", "assessment criterion")) for item in issues[:3])
        sentences.append(f"Management attention should prioritize {issue_names}.")
    while len(sentences) < 3:
        sentences.append("The assessment indicates that control evidence and management oversight should be strengthened in the highest-risk areas.")
    return " ".join(sentences[:5])


def _portfolio_business_summary(assessments: dict, findings: list[dict]) -> str:
    overall = _overall_score(assessments)
    risk = _highest_risk(assessments)
    compliant = partial = non_compliant = 0
    for _, data in _all_module_rows(assessments):
        matrix = data.get("evidence_matrix")
        if not isinstance(matrix, pd.DataFrame) or "Compliance Status" not in matrix:
            continue
        counts = matrix["Compliance Status"].value_counts()
        compliant += int(counts.get("Compliant", 0))
        partial += int(counts.get("Partially Compliant", 0))
        non_compliant += int(counts.get("Non-Compliant", 0))
    sentences = [
        f"The completed assessments produced an overall score of {overall}/100 with a {risk} portfolio risk rating.",
        f"The evidence review identified {compliant} compliant, {partial} partially compliant, and {non_compliant} non-compliant assessment criteria.",
    ]
    if findings:
        issue_names = ", ".join(str(item.get("Title", "assessment issue")) for item in findings[:3])
        sentences.append(f"Immediate management attention should focus on {issue_names}.")
    sentences.append("The recommended roadmap prioritizes control design, accountable ownership, operating evidence, and recurring management review.")
    return " ".join(sentences[:4])


def _evidence_line(item: dict) -> str:
    current_state = str(item.get("Current State", ""))
    marker = "Concise evidence references:"
    if marker.lower() in current_state.lower():
        start = current_state.lower().find(marker.lower()) + len(marker)
        evidence = current_state[start:].split(". Negative evidence", 1)[0].strip(" .")
        if evidence:
            return _limit_words(evidence, 18)
    documents = _top_documents_for_report(item.get("Required Documents", ""))
    return "; ".join(documents[:3]) if documents else "No sufficient evidence was identified."


def _top_documents_for_report(value: object) -> list[str]:
    items = []
    for part in re.split(r"[,;|]", str(value or "")):
        cleaned = re.sub(r"\s+", " ", part).strip(" .")
        if cleaned and _sentence_key(cleaned) not in {_sentence_key(item) for item in items}:
            items.append(cleaned)
        if len(items) == 5:
            break
    return items


def _add_top_findings_section(document: Document, guidance: list[dict]) -> None:
    document.add_heading("Top Findings", level=3)
    issues = _priority_guidance(guidance, limit=5)
    if not issues:
        document.add_paragraph("No material assessment findings were identified.")
        return
    for item in issues:
        paragraph = document.add_paragraph()
        finding = _clean_report_text(item.get("Gap Identified", item.get("Title", "Assessment criterion")), max_sentences=1, max_words=28)
        risk = _clean_report_text(item.get("Risk Impact", "Control weakness may affect operations and compliance."), max_sentences=1, max_words=22)
        for label, value in [("Finding", finding), ("Evidence", _evidence_line(item)), ("Risk Impact", risk)]:
            run = paragraph.add_run(f"{label}: ")
            run.bold = True
            paragraph.add_run(value)
            if label != "Risk Impact":
                paragraph.add_run("\n")


def _add_top_recommendations_section(document: Document, guidance: list[dict]) -> None:
    document.add_heading("Top Recommendations", level=3)
    issues = _priority_guidance(guidance, limit=5)
    if not issues:
        document.add_paragraph("No priority recommendations were generated.")
        return
    used_actions = set()
    for item in issues:
        title = str(item.get("Title", "Assessment criterion"))
        documents = _top_documents_for_report(item.get("Required Documents", ""))
        evidence_anchor = documents[0].lower() if documents else "documented control evidence"
        action = f"Formalize {title.lower()} through {evidence_anchor}, assign an accountable owner, and monitor exceptions to closure."
        action_key = _sentence_key(action)
        if not action or action_key in used_actions:
            action = f"Establish, approve, and monitor a documented {title.lower()} control with an accountable owner."
        used_actions.add(_sentence_key(action))
        evidence = "; ".join(documents[:3]) or "Approved control evidence and review records"
        paragraph = document.add_paragraph()
        values = [
            ("Issue", title),
            ("Recommended Action", action),
            ("Expected Evidence", evidence),
            ("Priority", str(item.get("Priority", "Medium"))),
            ("Timeline", str(item.get("Estimated Implementation Time", "3 Months"))),
        ]
        for index, (label, value) in enumerate(values):
            run = paragraph.add_run(f"{label}: ")
            run.bold = True
            paragraph.add_run(value)
            if index < len(values) - 1:
                paragraph.add_run("\n")


def _add_detailed_remediation_plan(document: Document, guidance: list[dict], scp_references: pd.DataFrame | None = None) -> None:
    used_actions = set()
    scp_reference = _top_scp_reference(scp_references)
    for item in _priority_guidance(guidance, limit=100):
        title = str(item.get("Title", "Assessment criterion"))
        document.add_heading(title, level=4)
        action_sentences = []
        for sentence in _split_report_sentences(item.get("Required Actions", "")):
            cleaned = _clean_report_text(sentence, max_sentences=1, max_words=42)
            key = _sentence_key(cleaned)
            if cleaned and key not in used_actions:
                used_actions.add(key)
                action_sentences.append(cleaned)
            if len(action_sentences) == 4:
                break
        if not action_sentences:
            action_sentences = [f"Implement and document the required {title.lower()} control, assign ownership, and monitor completion through management review."]
        for action in action_sentences:
            document.add_paragraph(action, style="List Number")
        documents = _top_documents_for_report(item.get("Required Documents", ""))
        _add_label_bullet(document, "Expected Evidence", "; ".join(documents[:5]) or "Approved procedure, operating records, and management review evidence")
        _add_label_bullet(document, "Priority", str(item.get("Priority", "Medium")))
        _add_label_bullet(document, "Timeline", str(item.get("Estimated Implementation Time", "3 Months")))
        _add_label_bullet(document, "Target Score", f"{item.get('Current Score', 0)}/5 -> {item.get('Target Score', 5)}/5")
        if scp_reference:
            _add_label_bullet(document, "SCP Reference", scp_reference)


def _portfolio_performance_grid(assessments: dict) -> pd.DataFrame:
    rows = []
    compliant_total = partial_total = non_compliant_total = 0
    for module_name, data in _all_module_rows(assessments):
        counts = {"Compliant": 0, "Partially Compliant": 0, "Non-Compliant": 0}
        matrix = data.get("evidence_matrix")
        if isinstance(matrix, pd.DataFrame) and "Compliance Status" in matrix:
            for status, count in matrix["Compliance Status"].value_counts().to_dict().items():
                counts[str(status)] = int(count)
        compliant_total += counts["Compliant"]
        partial_total += counts["Partially Compliant"]
        non_compliant_total += counts["Non-Compliant"]
        rows.append(
            {
                "Assessment Module": module_name.replace(" Assessment", ""),
                "Composite Score": f"{int(data.get('score', 0) or 0)} / 100",
                "Risk Exposure": str(data.get("risk_level", data.get("status", "Medium"))),
                "Controls Baseline": (
                    f"{counts['Compliant']} compliant, {counts['Partially Compliant']} partial, "
                    f"{counts['Non-Compliant']} non-compliant"
                ),
            }
        )
    if rows:
        rows.append(
            {
                "Assessment Module": "AGGREGATE ENTERPRISE PROFILE",
                "Composite Score": f"{_overall_score(assessments)} / 100",
                "Risk Exposure": _highest_risk(assessments),
                "Controls Baseline": (
                    f"{compliant_total} compliant, {partial_total} partial, "
                    f"{non_compliant_total} non-compliant"
                ),
            }
        )
    return pd.DataFrame(
        rows,
        columns=["Assessment Module", "Composite Score", "Risk Exposure", "Controls Baseline"],
    )


def _integrated_findings_matrix(assessments: dict, limit: int = 12) -> pd.DataFrame:
    rows = []
    for item in _summary_findings(assessments)[:limit]:
        actions = item.get("Recommendation Bullets", [])
        action = actions[0] if actions else _criterion_action(str(item.get("Title", "Assessment issue")))
        rows.append(
            {
                "Module": str(item.get("Module", "")).replace(" Assessment", ""),
                "Core Finding": str(item.get("Title", "Assessment issue")),
                "Risk Impact": str(item.get("Business Impact", "Control weakness may affect operations.")),
                "Targeted Remediation": _complete_sentence(action, 24),
                "Priority": str(item.get("Priority", "Medium")),
                "Timeline": str(item.get("Timeline", "90 Days")),
            }
        )
    return pd.DataFrame(
        rows,
        columns=["Module", "Core Finding", "Risk Impact", "Targeted Remediation", "Priority", "Timeline"],
    )


def _uploaded_document_register(documents: list[dict]) -> pd.DataFrame:
    rows = []
    for document in documents:
        if not isinstance(document, dict):
            continue
        rows.append(
            {
                "Document": document.get("name", document.get("filename", "Uploaded document")),
                "Type": document.get("type", document.get("file_type", "")),
            }
        )
    return pd.DataFrame(rows, columns=["Document", "Type"])


def build_word_report(client_profile: dict, documents: list[dict], assessments: dict) -> BytesIO:
    document = Document()
    _configure_document(document)
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Combined Detailed Assessment Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["navy"])
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("Integrated Capability, Compliance and Control Review").italic = True

    document.add_heading("Client Profile", level=1)
    for key, value in client_profile.items():
        label = key.replace("_", " ").title()
        document.add_paragraph(f"{label}: {value}")

    if not assessments:
        document.add_paragraph("No assessment modules completed.")
    else:
        findings = _summary_findings(assessments)
        document.add_heading("Enterprise Portfolio Overview", level=1)
        document.add_paragraph(_portfolio_business_summary(assessments, findings))

        document.add_heading("Unified Executive Performance Grid", level=2)
        _add_table(document, _portfolio_performance_grid(assessments), column_widths=[2.6, 1.0, 1.0, 2.9])

        document.add_heading("Integrated Findings and Remediation Matrix", level=2)
        _set_landscape(document)
        _add_table(
            document,
            _integrated_findings_matrix(assessments),
            column_widths=[1.4, 2.0, 2.6, 3.2, 0.9, 0.9],
        )
        _set_portrait(document)

        document.add_heading("Enterprise 30-60-90 Day Roadmap", level=2)
        _add_roadmap_summary(document, findings)

    document.add_heading("Detailed Module Assessments", level=1)

    for module_name, data in assessments.items():
        report_guidance = _deduplicate_report_guidance(data.get("full_compliance_guidance") or [])
        _set_portrait(document)
        if document.paragraphs:
            document.add_page_break()
        document.add_heading(module_name, level=2)
        document.add_paragraph(f"Assessment Basis: {_assessment_basis_label(data.get('assessment_mode'))}")
        if "esg_score" in data and "climate_score" in data:
            document.add_paragraph(f"ESG Findings Score: {data['esg_score']} / 100")
            document.add_paragraph(f"Climate Compliance Findings Score: {data['climate_score']} / 100")
            if "sustainability_dashboard" in data:
                document.add_heading("Combined Sustainability Dashboard", level=3)
                _add_table(document, data["sustainability_dashboard"])
            sustainability_kpis = data.get("combined_sustainability_kpi_dashboard")
            if isinstance(sustainability_kpis, pd.DataFrame) and not sustainability_kpis.empty:
                document.add_heading("Combined Sustainability KPI Framework", level=3)
                columns = [
                    column
                    for column in [
                        "Sustainability Section",
                        "KPI Name",
                        "Current",
                        "Current Value Reason",
                        "Target",
                        "Maturity",
                        "Risk",
                        "Status",
                    ]
                    if column in sustainability_kpis.columns
                ]
                _add_table(document, sustainability_kpis[columns])

        document.add_heading("Module Overview", level=3)
        _add_executive_scorecard(document, module_name, data)
        document.add_paragraph(_business_summary(module_name, data, report_guidance))
        _add_critical_issues_at_glance(document, report_guidance)

        if "scoring_model" in data:
            _add_risk_heat_map(document, data["scoring_model"])

        if "evidence_matrix" in data:
            _add_compliance_dashboard(document, data["evidence_matrix"])

        if "scoring_model" in data:
            _set_landscape(document)
            document.add_heading("Assessment Scoring Model", level=3)
            _add_table(
                document,
                _compact_scoring_model(data["scoring_model"]),
                column_widths=[2.4, 0.8, 0.8, 1.0, 4.4, 1.0],
            )
            _set_portrait(document)

        if report_guidance:
            document.add_heading("Detailed Remediation Plan", level=3)
            _add_detailed_remediation_plan(document, report_guidance, data.get("scp_references"))

        if "evidence_matrix" in data:
            document.add_heading("Evidence Mapping Matrix", level=3)
            _add_table(
                document,
                _evidence_summary_table(data.get("evidence_matrix"), data.get("evidence_register")),
                column_widths=[2.6, 2.0, 2.0],
            )

        if "scp_references" in data and not data["scp_references"].empty:
            document.add_heading("Knowledge Base Used", level=3)
            _add_table(document, _compact_references(data["scp_references"]))

        if data.get("procurement_radar") is not None:
            document.add_heading("Maturity Radar", level=3)
            _add_table(document, data["procurement_radar"])
        elif "scoring_model" in data:
            document.add_heading("Maturity Radar", level=3)
            _add_table(document, _maturity_radar_from_scoring_model(data["scoring_model"]))

        if "kpi_table" in data:
            document.add_heading("KPI Framework", level=3)
            _add_table(document, _compact_kpi_table(data["kpi_table"]), column_widths=[1.6, 2.2, 1.2, 0.8, 1.1])
            document.add_heading("Appendix KPI Formulas", level=3)
            _add_table(document, _kpi_formula_appendix(data["kpi_table"]), column_widths=[2.0, 3.3, 2.0])

    document.add_page_break()
    document.add_heading("Appendix: Uploaded Document Register", level=1)
    uploaded_register = _uploaded_document_register(documents)
    if uploaded_register.empty:
        document.add_paragraph("No uploaded documents were recorded for this assessment portfolio.")
    else:
        _add_table(document, uploaded_register, column_widths=[5.4, 1.1])

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_combined_detailed_report_docx(client_profile: dict, documents: list[dict], assessments: dict) -> BytesIO:
    return build_word_report(client_profile, documents, assessments)


def _add_list(document: Document, items: list[str]) -> None:
    if not items:
        document.add_paragraph("No items available.")
        return

    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _add_full_compliance_guidance(document: Document, guidance: list[dict]) -> None:
    for item in guidance:
        title = item.get("Title", "Assessment area")
        item_type = item.get("Type", "Assessment Criterion")
        document.add_heading(f"{title} ({item_type})", level=4)
        for label in [
            "Current State",
            "Gap Identified",
            "Risk Impact",
            "Required Actions",
            "Required Documents",
            "Required Controls",
            "Evidence Expected",
            "Target State",
            "Expected Score Improvement",
            ]:
            document.add_paragraph(f"{label}:")
            document.add_paragraph(str(item.get(label, "Not available.")))


def _score_color(score: int) -> str:
    if score <= 20:
        return TRAFFIC_COLORS["red"]
    if score <= 40:
        return TRAFFIC_COLORS["orange"]
    if score <= 60:
        return TRAFFIC_COLORS["amber"]
    if score <= 80:
        return TRAFFIC_COLORS["light_green"]
    return TRAFFIC_COLORS["green"]


def _risk_color(risk: str) -> str:
    return {
        "Critical": TRAFFIC_COLORS["dark_red"],
        "High": TRAFFIC_COLORS["red"],
        "Medium": TRAFFIC_COLORS["amber"],
        "Low": TRAFFIC_COLORS["green"],
    }.get(str(risk), TRAFFIC_COLORS["amber"])


def _status_color(status: str) -> str:
    return {
        "Compliant": TRAFFIC_COLORS["green"],
        "Partially Compliant": TRAFFIC_COLORS["amber"],
        "Non-Compliant": TRAFFIC_COLORS["red"],
        "Green": TRAFFIC_COLORS["green"],
        "Amber": TRAFFIC_COLORS["amber"],
        "Red": TRAFFIC_COLORS["red"],
    }.get(str(status), TRAFFIC_COLORS["amber"])


def _roadmap_color(priority: str) -> str:
    return {
        "Critical": TRAFFIC_COLORS["red"],
        "High": TRAFFIC_COLORS["orange"],
        "Medium": TRAFFIC_COLORS["yellow"],
        "Low": TRAFFIC_COLORS["green"],
    }.get(str(priority), TRAFFIC_COLORS["yellow"])


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell(cell, text: str, bold: bool = False, fill: str | None = None, color: str | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for no_wrap in tc_pr.findall(qn("w:noWrap")):
        tc_pr.remove(no_wrap)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if fill:
        _shade_cell(cell, fill)


def _format_summary_table(table) -> None:
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_autofit(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
            if row_index == 0:
                _shade_cell(cell, TRAFFIC_COLORS["navy"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["white"])
        if row_index == 0:
            _repeat_table_header(row)


def _add_summary_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["navy"])


def _all_module_rows(assessments: dict) -> list[tuple[str, dict]]:
    return [(str(module), data) for module, data in assessments.items() if isinstance(data, dict)]


def _overall_score(assessments: dict) -> int:
    scores = [int(data.get("score", 0) or 0) for _, data in _all_module_rows(assessments)]
    return round(sum(scores) / len(scores)) if scores else 0


def _highest_risk(assessments: dict) -> str:
    risks = [str(data.get("risk_level", data.get("status", "Medium"))) for _, data in _all_module_rows(assessments)]
    ordered = sorted(risks, key=lambda item: PRIORITY_ORDER.get(item, 2))
    return ordered[0] if ordered else "Medium"


def _compliance_percent(assessments: dict) -> int:
    compliant = 0
    total = 0
    for _, data in _all_module_rows(assessments):
        matrix = data.get("evidence_matrix")
        if isinstance(matrix, pd.DataFrame) and not matrix.empty and "Compliance Status" in matrix:
            total += len(matrix)
            compliant += int((matrix["Compliance Status"] == "Compliant").sum())
    return round((compliant / total) * 100) if total else _overall_score(assessments)


def _summary_findings(assessments: dict) -> list[dict]:
    findings = []
    for module_name, data in _all_module_rows(assessments):
        guidance = _deduplicate_report_guidance(data.get("full_compliance_guidance") or [])
        module_findings_before = len(findings)
        for item in guidance:
            if not isinstance(item, dict) or item.get("Type") == "KPI Framework":
                continue
            current = _score_value(item.get("Current Score", 0))
            target = _score_value(item.get("Target Score", 5))
            if current >= target:
                continue
            priority = _normalize_priority(item.get("Priority", "Medium"))
            title = _clean_line(item.get("Title") or item.get("Criteria") or "Assessment gap", 9)
            findings.append(
                {
                    "Module": module_name,
                    "Title": title,
                    "Finding": _clean_line(f"{module_name}: {title}", 14),
                    "Score": f"{current}/5",
                    "Priority": priority,
                    "Risk": priority,
                    "Missing Evidence": _missing_evidence_for_criterion(data, title) or _missing_evidence(item),
                    "Recommendation Bullets": _recommendation_bullets(item),
                    "Business Impact": _business_impact(item, module_name, priority),
                    "Timeline": {"Critical": "30 Days", "High": "60 Days", "Medium": "90 Days", "Low": "90 Days"}[priority],
                }
            )
        if len(findings) > module_findings_before or guidance:
            continue
        scoring = data.get("scoring_model")
        if not isinstance(scoring, pd.DataFrame) or scoring.empty:
            continue
        for _, row in scoring.iterrows():
            current = _score_value(row.get("Score 0-5", 0))
            if current >= 5:
                continue
            title = _clean_line(row.get("Assessment Criteria", "Assessment gap"), 9)
            priority = _normalize_priority(row.get("Risk Rating", "Medium"))
            findings.append(
                {
                    "Module": module_name,
                    "Title": title,
                    "Finding": _clean_line(f"{module_name}: {title}", 14),
                    "Score": f"{current}/5",
                    "Priority": priority,
                    "Risk": priority,
                    "Missing Evidence": _missing_evidence_for_criterion(data, title),
                    "Recommendation Bullets": _split_actions(row.get("Recommended Improvement Action", ""))[:3]
                    or ["Define the missing control", "Assign an accountable owner", "Track closure evidence"],
                    "Business Impact": _business_impact(row.to_dict(), module_name, priority),
                    "Timeline": _timeline_label("1 Month" if priority in {"Critical", "High"} else "3 Months"),
                }
            )
    unique_findings = []
    seen = set()
    for finding in findings:
        key = (str(finding.get("Module", "")), _sentence_key(finding.get("Title", "")))
        if key in seen:
            continue
        seen.add(key)
        finding["Recommendation Bullets"] = list(dict.fromkeys(finding.get("Recommendation Bullets", [])))[:3]
        unique_findings.append(finding)
    return sorted(
        unique_findings,
        key=lambda item: (
            PRIORITY_ORDER.get(str(item.get("Priority", "Medium")), 2),
            _score_value(str(item.get("Score", "0")).split("/")[0]),
            str(item.get("Module", "")),
        ),
    )


def _score_value(value: object) -> int:
    try:
        return max(0, min(5, int(float(str(value).strip()))))
    except (TypeError, ValueError):
        return 0


def _normalize_priority(value: object) -> str:
    text = str(value or "Medium").strip().title()
    return text if text in PRIORITY_ORDER else "Medium"


def _clean_line(value: object, max_words: int = 12) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" .;-")
    return _limit_words(text, max_words)


def _split_actions(value: object) -> list[str]:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if not text:
        return []
    parts = re.split(r"(?:\n|;|\u2022|\- )+|(?<=[.!?])\s+", text)
    actions = []
    for part in parts:
        action = _complete_sentence(part, 24)
        if action and _sentence_key(action) not in {_sentence_key(item) for item in actions}:
            actions.append(action)
    return actions


def _complete_sentence(value: object, max_words: int = 24) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" .;-:")
    if not text:
        return ""
    words = text.split()
    if len(words) > max_words:
        text = " ".join(words[:max_words]).rstrip(" ,;:")
    incomplete_endings = {
        "that", "for", "before", "before the", "and", "or", "to", "with", "through",
        "the process", "before the process", "management review before the process",
    }
    lowered = text.lower().rstrip(".?!")
    if any(lowered == ending or lowered.endswith(f" {ending}") for ending in incomplete_endings):
        return ""
    if text[-1] not in ".!?":
        text += "."
    return text


def _missing_evidence(item: dict, limit: int = 3) -> list[str]:
    evidence_text = item.get("Required Documents") or item.get("Required Evidence") or item.get("Evidence Required") or ""
    docs = _split_actions(evidence_text)
    if not docs:
        docs = ["Approved evidence record"]
    return docs[:limit]


def _missing_evidence_for_criterion(data: dict, criterion: str, limit: int = 3) -> list[str]:
    matrix = data.get("evidence_matrix")
    if not isinstance(matrix, pd.DataFrame) or matrix.empty:
        return []
    if "Assessment Criteria" not in matrix.columns:
        return []
    criterion_text = str(criterion).lower()
    matches = matrix[matrix["Assessment Criteria"].astype(str).str.lower().str.contains(re.escape(criterion_text[:30]), na=False)]
    if matches.empty:
        criterion_tokens = {token for token in re.findall(r"[a-z]{4,}", criterion_text)}
        if criterion_tokens:
            matches = matrix[
                matrix["Assessment Criteria"]
                .astype(str)
                .str.lower()
                .apply(lambda value: bool(criterion_tokens.intersection(set(re.findall(r"[a-z]{4,}", value)))))
            ]
    if matches.empty:
        return []
    row = matches.iloc[0]
    explicit_missing = row.get("Missing Evidence", "")
    if explicit_missing and str(explicit_missing).strip().lower() not in {"nan", "none"}:
        return _split_actions(explicit_missing)[:limit]
    required = row.get("Required Evidence", "")
    missing = _missing_required_labels(required, row)
    if missing:
        return [_clean_line(item, 6) for item in missing[:limit]]
    return _split_actions(required)[:limit]


def _recommendation_bullets(item: dict, limit: int = 3) -> list[str]:
    title = str(item.get("Title") or item.get("Criteria") or "control gap")
    specific_action = _criterion_action(title)
    if specific_action:
        return [specific_action]
    actions = _split_actions(item.get("Required Actions") or item.get("Recommendation") or item.get("Recommended Actions"))
    return (actions or [f"Approve a documented {title.lower()} control, assign an accountable owner, and review operating evidence each reporting period."])[:limit]


def _criterion_action(title: str) -> str:
    text = title.lower()
    actions = [
        (("strategy", "operating model"), "Approve a procurement strategy and operating model with defined objectives, decision rights, governance forums, and accountable process owners."),
        (("purchase-to-pay", "accounts payable", "invoice"), "Enforce no-PO-no-pay and three-way matching, segregate invoice approval, and document all payment exceptions."),
        (("policy", "procedures"), "Approve and publish the governing policy, authority thresholds, and procedures, then train users and monitor exceptions."),
        (("category",), "Approve category plans with demand forecasts, sourcing milestones, savings targets, risk actions, and named business owners."),
        (("strategic sourcing", "tender"), "Mandate competitive sourcing thresholds, documented evaluations, conflict declarations, award approvals, and a complete tender audit trail."),
        (("due diligence", "onboarding"), "Require legal, tax, sanctions, financial, ESG, and conflict checks before supplier activation and retain approval evidence."),
        (("supplier performance",), "Implement supplier scorecards, scheduled performance reviews, corrective action plans, and escalation for repeated service failures."),
        (("risk and compliance", "compliance monitoring"), "Maintain a procurement risk register, test policy compliance, assign remediation owners, and report overdue actions to management."),
        (("spend visibility", "analytics"), "Consolidate spend data and issue recurring category, supplier, savings, and compliance reports with documented management review."),
        (("technology", "automation"), "Implement controlled digital workflows for requisitions, approvals, sourcing, contracting, and supplier management with access reviews."),
        (("contract",), "Maintain a complete contract register with ownership, obligations, renewal alerts, performance reviews, and documented close-out controls."),
        (("climate", "emission", "ghg"), "Assign climate data owners, document calculation controls, verify source evidence, and approve recurring emissions reporting."),
        (("esg", "sustainability"), "Approve ESG governance, assign accountable data owners, monitor material targets, and retain evidence supporting external disclosures."),
        (("bank", "cash"), "Complete and independently review bank reconciliations each month, investigate aged items, and retain approval evidence."),
        (("financial reporting", "reconciliation"), "Complete account reconciliations to an approved timetable, investigate variances, and evidence independent management review."),
    ]
    for terms, action in actions:
        if any(term in text for term in terms):
            return action
    return ""


def _business_impact(item: dict, module_name: str, priority: str) -> str:
    text = " ".join([str(item.get("Title", "")), str(item.get("Criteria", ""))]).lower()
    if "purchase-to-pay" in text or "invoice" in text:
        impact = "Unauthorized or duplicate payments may cause direct financial loss."
    elif "policy" in text or "procedure" in text:
        impact = "Inconsistent approvals may create control breaches and audit findings."
    elif "category" in text:
        impact = "Unplanned demand and fragmented sourcing may increase total cost."
    elif "strategic sourcing" in text or "tender" in text:
        impact = "Weak competition may reduce value and undermine award transparency."
    elif "due diligence" in text or "onboarding" in text:
        impact = "Unscreened suppliers may expose the organization to fraud or sanctions."
    elif "supplier performance" in text:
        impact = "Unmanaged supplier failures may disrupt service and contractual outcomes."
    elif "spend" in text or "analytics" in text:
        impact = "Poor spend visibility may conceal leakage and sourcing opportunities."
    elif "technology" in text or "automation" in text:
        impact = "Manual workflows may increase errors, delays, and unauthorized access."
    elif "risk" in text or "compliance" in text:
        impact = "Unmonitored exceptions may remain unresolved and recur across transactions."
    elif "contract" in text:
        impact = "Missed obligations and renewals may cause service disruption and value leakage."
    elif "climate" in text or "emission" in text or "ghg" in text:
        impact = "Incomplete climate data may cause regulatory and disclosure failures."
    elif "esg" in text or "sustainability" in text:
        impact = "Weak ESG controls may impair disclosures and stakeholder confidence."
    elif "bank" in text or "cash" in text:
        impact = "Unreconciled cash balances may conceal error, fraud, or misstatement."
    elif "finance" in module_name.lower() or "audit" in module_name.lower():
        impact = "Control deficiencies may cause financial misstatement and repeat audit findings."
    elif priority == "Critical":
        impact = "Material control weakness may affect compliance and operations."
    else:
        impact = "Control gap may reduce assurance and management visibility."
    return _limit_words(impact, 15)


KPI_EVIDENCE_ALIASES = {
    "Cost savings %": ["cost savings", "realized savings", "savings rate"],
    "PO cycle time": ["po cycle time", "purchase order cycle time", "requisition to po cycle time"],
    "Spend under contract %": ["spend under contract", "contracted spend"],
    "Supplier on-time delivery %": ["supplier on-time delivery", "supplier on time delivery", "on-time delivery", "otd"],
    "Maverick spend %": ["maverick spend", "off-contract spend"],
    "Competitive sourcing coverage %": ["competitive sourcing coverage", "competitively sourced spend"],
    "Supplier evaluation completion %": ["supplier evaluation completion", "evaluated suppliers"],
    "Contract renewal compliance %": ["contract renewal compliance", "on-time renewals"],
    "Procurement policy compliance %": ["procurement policy compliance", "compliant procurement transactions"],
    "Supplier risk review completion %": ["supplier risk review completion", "reviewed high-risk suppliers"],
}

KPI_NON_CURRENT_TERMS = {
    "target",
    "benchmark",
    "goal",
    "threshold",
    "assessment score",
    "maturity score",
    "compliance score",
    "overall score",
}


def _kpi_evidence_records(data: dict) -> list[dict]:
    records: list[dict] = []
    evidence_register = data.get("evidence_register")
    if isinstance(evidence_register, pd.DataFrame) and not evidence_register.empty:
        for _, row in evidence_register.iterrows():
            extract = str(row.get("Extract", "")).strip()
            if extract:
                records.append(
                    {
                        "source": str(row.get("Source Document", "Evidence register")),
                        "text": extract,
                    }
                )

    for document in data.get("extracted_documents", []) or []:
        if not isinstance(document, dict):
            continue
        source = str(document.get("name", document.get("filename", "Uploaded document")))
        text = str(document.get("text", document.get("content", ""))).strip()
        if text:
            records.append({"source": source, "text": text})
    return records


def _structured_kpi_measurement(kpi_name: str, data: dict) -> dict | None:
    for key in ("calculated_kpi_records", "kpi_records", "kpi_results"):
        records = data.get(key)
        if isinstance(records, pd.DataFrame):
            iterable = records.to_dict("records")
        elif isinstance(records, list):
            iterable = records
        else:
            continue
        for record in iterable:
            if not isinstance(record, dict):
                continue
            record_name = str(record.get("KPI Name", record.get("KPI", record.get("name", "")))).strip()
            if record_name.casefold() != kpi_name.casefold():
                continue
            raw_value = record.get("Current Value", record.get("Current", record.get("value")))
            numeric = re.search(r"-?\d+(?:\.\d+)?", str(raw_value or ""))
            if not numeric:
                continue
            return {
                "value": float(numeric.group()),
                "evidence_count": int(record.get("Evidence Count", 1) or 1),
                "data_source": str(record.get("Data Source", record.get("Source", "Calculated KPI record"))),
                "calculation_method": str(record.get("Calculation Method", "Calculated KPI record")),
                "maturity_score": record.get("Maturity Score", record.get("Maturity")),
            }
    return None


def _extract_explicit_kpi_values(kpi_name: str, data: dict, unit: str) -> list[dict]:
    aliases = KPI_EVIDENCE_ALIASES.get(kpi_name, [kpi_name.lower().replace("%", "").strip()])
    value_pattern = r"(?<![\d.])(\d{1,3}(?:\.\d+)?)\s*%" if unit == "%" else r"(?<![\d.])(\d{1,3}(?:\.\d+)?)\s*days?\b"
    matches: list[dict] = []
    for record in _kpi_evidence_records(data):
        segments = [segment.strip() for segment in re.split(r"[\r\n]+|(?<=[.!?;])\s+", record["text"]) if segment.strip()]
        for segment in segments:
            lowered = segment.lower()
            alias = next((candidate for candidate in aliases if candidate in lowered), None)
            if not alias:
                continue
            if any(term in lowered for term in KPI_NON_CURRENT_TERMS):
                continue
            numeric_matches = list(re.finditer(value_pattern, segment, flags=re.IGNORECASE))
            if not numeric_matches:
                continue
            alias_start = lowered.find(alias)
            numeric = min(numeric_matches, key=lambda match: abs(match.start() - alias_start))
            matches.append(
                {
                    "value": float(numeric.group(1)),
                    "source": record["source"],
                    "extract": segment,
                }
            )
    return matches


def _kpi_metadata(kpi_name: str, target: object) -> dict:
    if kpi_name in KPI_TARGET_METADATA:
        return KPI_TARGET_METADATA[kpi_name].copy()
    target_text = str(target or "")
    name_text = str(kpi_name or "")
    numeric = re.search(r"(\d{1,3}(?:\.\d+)?)", target_text)
    lowered = f"{name_text} {target_text}".lower()
    unit = "days" if "day" in lowered else "%" if "%" in lowered or " x 100" in lowered else ""
    lower_terms = ["less", "reduction", "cycle time", "leakage", "exception rate", "intensity", "overdue ratio"]
    direction = "lower_is_better" if any(term in lowered for term in lower_terms) else "higher_is_better"
    return {
        "unit": unit,
        "target_value": float(numeric.group(1)) if numeric else None,
        "target_direction": direction,
    }


def _kpi_target_display(kpi_name: str, target: object = "") -> str:
    metadata = _kpi_metadata(kpi_name, target)
    return format_kpi_target(metadata["target_value"], metadata["unit"])


def _format_kpi_value(value: float | int | None, unit: str) -> str:
    if value is None or pd.isna(value):
        return "Data Not Available"
    number = int(value) if float(value).is_integer() else round(float(value), 1)
    if unit:
        return f"{number}{unit}" if unit == "%" else f"{number} {unit}"
    return str(number)


def _kpi_measurement(kpi_name: str, data: dict, unit: str, configured_source: str) -> dict:
    structured = _structured_kpi_measurement(kpi_name, data)
    if structured:
        return structured

    matches = _extract_explicit_kpi_values(kpi_name, data, unit)
    if not matches:
        aliases = KPI_EVIDENCE_ALIASES.get(kpi_name, [kpi_name.lower().replace("%", "").strip()])
        framework_sources = [
            record["source"]
            for record in _kpi_evidence_records(data)
            if any(alias in record["text"].lower() for alias in aliases)
        ]
        framework_evidence_count = len(set(framework_sources))
        return {
            "value": None,
            "evidence_count": framework_evidence_count,
            "data_source": "; ".join(dict.fromkeys(framework_sources)) if framework_sources else configured_source,
            "calculation_method": "KPI referenced without sufficient numerical data" if framework_evidence_count else "No KPI framework evidence identified",
            "maturity_score": 1 if framework_evidence_count else 0,
        }

    value = matches[0]["value"]
    return {
        "value": max(0, min(100, value)) if unit == "%" else max(0, value),
        "evidence_count": len(matches),
        "data_source": "; ".join(dict.fromkeys(match["source"] for match in matches)),
        "calculation_method": "Explicit KPI value extracted from uploaded evidence",
        "maturity_score": 3 if len(matches) > 1 else 2,
    }


def _kpi_current_from_evidence(kpi_name: str, data: dict, unit: str) -> float | None:
    return _kpi_measurement(kpi_name, data, unit, "Uploaded evidence")["value"]


def _kpi_assessment_rows(assessments: dict) -> list[tuple[str, dict]]:
    rows = []
    for module_name, data in _all_module_rows(assessments):
        if isinstance(data.get("kpi_table"), pd.DataFrame):
            rows.append((module_name, data))
            continue
        for section_name, key in [
            ("ESG Governance & Reporting", "esg_assessment"),
            ("UAE Climate Law Compliance", "climate_assessment"),
        ]:
            child = data.get(key)
            if isinstance(child, dict) and isinstance(child.get("kpi_table"), pd.DataFrame):
                rows.append((f"{module_name} - {section_name}", child))
    return rows


def _coerce_kpi_maturity(value: object, evidence_count: int, measured: bool) -> int:
    if value is not None:
        numeric = re.search(r"[0-5]", str(value))
        if numeric:
            return int(numeric.group())
    if not measured:
        return 1 if evidence_count else 0
    return 3 if evidence_count > 1 else 2


def _summary_kpis(assessments: dict, limit: int | None = 10) -> pd.DataFrame:
    rows = []
    for module_name, data in _kpi_assessment_rows(assessments):
        table = data.get("kpi_table")
        if not isinstance(table, pd.DataFrame) or table.empty:
            continue
        for _, row in table.iterrows():
            name = str(row.get("KPI Name", ""))
            metadata = _kpi_metadata(name, row.get("Target", ""))
            unit = str(metadata["unit"])
            configured_source = str(row.get("Data Source", "Configured KPI source"))
            measurement = _kpi_measurement(name, data, unit, configured_source)
            current = measurement["value"]
            target_value = metadata["target_value"]
            target = None if target_value is None or pd.isna(target_value) else float(target_value)
            direction = str(metadata["target_direction"])
            if current is not None:
                status = "Measured"
            elif int(measurement["evidence_count"]):
                status = "Data Not Available"
            else:
                status = "Not Established"
            maturity_score = _coerce_kpi_maturity(
                measurement.get("maturity_score"),
                int(measurement["evidence_count"]),
                current is not None,
            )
            if current is not None:
                current_reason = (
                    f"Measured from {measurement['data_source']} using "
                    f"{str(measurement['calculation_method']).lower()}."
                )
            elif int(measurement["evidence_count"]):
                current_reason = (
                    "KPI-related evidence was found, but the uploaded records do not contain sufficient numerical data "
                    "to calculate the current value."
                )
            else:
                current_reason = "No KPI framework or KPI-specific evidence was identified in the uploaded documents."
            rows.append(
                {
                    "Module": module_name,
                    "KPI": name,
                    "Unit": unit,
                    "Target Direction": direction,
                    "Current": current,
                    "Current Display": (
                        _format_kpi_value(current, unit)
                        if current is not None
                        else ("Data Not Available" if int(measurement["evidence_count"]) else "Not Established")
                    ),
                    "Current Value Reason": current_reason,
                    "Target": target,
                    "Target Display": _kpi_target_display(name, row.get("Target", "")),
                    "Status": status,
                    "Maturity Score": maturity_score,
                    "Maturity": format_kpi_maturity(maturity_score),
                    "Risk": kpi_risk_from_maturity(maturity_score),
                    "Owner": row.get("Owner", "Function Owner"),
                    "Priority": _kpi_priority(status),
                    "Data Source": measurement["data_source"],
                    "Evidence Count": measurement["evidence_count"],
                    "Calculation Method": measurement["calculation_method"],
                }
            )
    selected_rows = rows
    if limit is not None:
        grouped: dict[str, list[dict]] = {}
        for row in rows:
            grouped.setdefault(str(row["Module"]), []).append(row)
        selected_rows = []
        position = 0
        while len(selected_rows) < limit and any(position < len(items) for items in grouped.values()):
            for items in grouped.values():
                if position < len(items):
                    selected_rows.append(items[position])
                    if len(selected_rows) == limit:
                        break
            position += 1
    dataframe = pd.DataFrame(
        selected_rows,
        columns=[
            "Module",
            "KPI",
            "Unit",
            "Target Direction",
            "Current",
            "Current Display",
            "Current Value Reason",
            "Target",
            "Target Display",
            "Status",
            "Maturity Score",
            "Maturity",
            "Risk",
            "Owner",
            "Priority",
            "Data Source",
            "Evidence Count",
            "Calculation Method",
        ],
    )
    dataframe.attrs["placeholder_warning"] = _kpi_placeholder_warning(dataframe)
    return dataframe


def build_kpi_dashboard_dataframe(assessments: dict) -> pd.DataFrame:
    dataframe = _summary_kpis(assessments, limit=None).copy()
    established_modules = set(dataframe["Module"]) if not dataframe.empty else set()
    missing_rows = []
    for module_name in assessments:
        if module_name in established_modules:
            continue
        missing_rows.append(
            {
                "Module": module_name,
                "KPI": "KPI Framework",
                "Current Display": "Not Established",
                "Current Value Reason": "No KPI framework or KPI-specific evidence was identified in the uploaded documents.",
                "Target Display": "Not Established",
                "Maturity Score": 0,
                "Maturity": format_kpi_maturity(0),
                "Risk": kpi_risk_from_maturity(0),
                "Status": "Not Established",
            }
        )
    if missing_rows:
        dataframe = pd.concat([dataframe, pd.DataFrame(missing_rows)], ignore_index=True)
    return dataframe


def _kpi_placeholder_warning(kpis: pd.DataFrame) -> str:
    if not isinstance(kpis, pd.DataFrame) or kpis.empty or "Current" not in kpis.columns:
        return ""
    measured = kpis["Current"].dropna()
    if len(measured) >= 2 and len(measured) == len(kpis) and measured.nunique() == 1:
        return "Possible placeholder KPI values detected."
    return ""


def build_kpi_diagnostics(assessment: dict) -> pd.DataFrame:
    diagnostics = _summary_kpis({"Assessment": assessment})
    columns = ["KPI", "Data Source", "Evidence Count", "Calculation Method"]
    result = diagnostics[columns].rename(columns={"KPI": "KPI Name"}) if not diagnostics.empty else pd.DataFrame(columns=["KPI Name", "Data Source", "Evidence Count", "Calculation Method"])
    result.attrs["placeholder_warning"] = diagnostics.attrs.get("placeholder_warning", "")
    return result


def _kpi_status(current: float | None, target: float | None, direction: str) -> str:
    if current is None or target is None:
        return "Data Not Available"
    if direction == "lower_is_better":
        if current <= target:
            return "Green"
        if current <= target * 1.2:
            return "Amber"
        return "Red"
    if current >= target:
        return "Green"
    if current >= target * 0.8:
        return "Amber"
    return "Red"


def _kpi_priority(status: str) -> str:
    return {"Red": "High", "Amber": "Medium", "Green": "Low"}.get(str(status), "Medium")


def _maturity_rows(assessments: dict) -> list[tuple[str, int, int]]:
    buckets = {
        "Governance": [],
        "Strategy": [],
        "Risk": [],
        "Compliance": [],
        "Digital": [],
        "Performance": [],
    }
    for _, data in _all_module_rows(assessments):
        scoring = data.get("scoring_model")
        if not isinstance(scoring, pd.DataFrame) or scoring.empty:
            continue
        for _, row in scoring.iterrows():
            criterion = str(row.get("Assessment Criteria", "")).lower()
            score = int(row.get("Score 0-5", 0) or 0) * 20
            if any(term in criterion for term in ["governance", "policy", "accountability"]):
                buckets["Governance"].append(score)
            elif any(term in criterion for term in ["strategy", "materiality", "category", "planning"]):
                buckets["Strategy"].append(score)
            elif "risk" in criterion or "due diligence" in criterion:
                buckets["Risk"].append(score)
            elif any(term in criterion for term in ["compliance", "reporting", "controls", "audit"]):
                buckets["Compliance"].append(score)
            elif any(term in criterion for term in ["technology", "data", "digital", "automation"]):
                buckets["Digital"].append(score)
            else:
                buckets["Performance"].append(score)
    return [
        (name, round(sum(values) / len(values)) if values else _overall_score(assessments), 95)
        for name, values in buckets.items()
    ]


def _add_card_row(document: Document, cards: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=len(cards))
    table.autofit = True
    for index, (label, value, fill) in enumerate(cards):
        cell = table.cell(0, index)
        _shade_cell(cell, fill)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["white"])
        p2 = cell.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r2 = p2.add_run(value)
        r2.bold = True
        r2.font.name = "Arial"
        r2.font.size = Pt(14)
        r2.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["white"])


def _add_top_list_table(document: Document, findings: list[dict]) -> None:
    table = document.add_table(rows=1, cols=2)
    _set_cell(table.cell(0, 0), "Top 5 Findings", bold=True, color=TRAFFIC_COLORS["white"])
    _set_cell(table.cell(0, 1), "Top 5 Recommendations", bold=True, color=TRAFFIC_COLORS["white"])
    row = table.add_row().cells
    row[0].text = ""
    row[1].text = ""
    for idx, item in enumerate(findings[:5]):
        p = row[0].paragraphs[0] if idx == 0 else row[0].add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"• {item.get('Finding', '')}").font.size = Pt(8)
        p2 = row[1].paragraphs[0] if idx == 0 else row[1].add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        recommendation = str(item.get("Recommendation", ""))
        p2.add_run(f"• {_limit_words(recommendation, 16)}").font.size = Pt(8)
    _format_summary_table(table)


def _add_scorecard_table(document: Document, assessments: dict) -> None:
    table = document.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Module", "Score", "Risk", "Mode"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for module_name, data in _all_module_rows(assessments):
        score = int(data.get("score", 0) or 0)
        risk = str(data.get("risk_level", data.get("status", "Medium")))
        cells = table.add_row().cells
        _set_cell(cells[0], module_name)
        _set_cell(cells[1], f"{score}/100", bold=True, fill=_score_color(score))
        _set_cell(cells[2], risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
        _set_cell(cells[3], data.get("assessment_mode", "Not specified"))
    _format_summary_table(table)


def _add_risk_heat_table(document: Document, findings: list[dict]) -> None:
    counts = {risk: 0 for risk in ["Critical", "High", "Medium", "Low"]}
    for item in findings:
        priority = str(item.get("Priority", "Medium"))
        counts[priority] = counts.get(priority, 0) + 1
    table = document.add_table(rows=1, cols=4)
    for idx, risk in enumerate(["Critical", "High", "Medium", "Low"]):
        _set_cell(table.cell(0, idx), risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
    row = table.add_row().cells
    for idx, risk in enumerate(["Critical", "High", "Medium", "Low"]):
        _set_cell(row[idx], str(counts.get(risk, 0)), bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])


def _add_compliance_summary_table(document: Document, assessments: dict) -> None:
    counts = {"Compliant": 0, "Partially Compliant": 0, "Non-Compliant": 0}
    for _, data in _all_module_rows(assessments):
        matrix = data.get("evidence_matrix")
        if isinstance(matrix, pd.DataFrame) and not matrix.empty and "Compliance Status" in matrix:
            for status, count in matrix["Compliance Status"].value_counts().to_dict().items():
                counts[str(status)] = counts.get(str(status), 0) + int(count)
    table = document.add_table(rows=1, cols=3)
    for idx, status in enumerate(["Compliant", "Partially Compliant", "Non-Compliant"]):
        _set_cell(table.cell(0, idx), status, bold=True, fill=_status_color(status), color=TRAFFIC_COLORS["white"] if status != "Partially Compliant" else TRAFFIC_COLORS["black"])
    row = table.add_row().cells
    for idx, status in enumerate(["Compliant", "Partially Compliant", "Non-Compliant"]):
        _set_cell(row[idx], str(counts.get(status, 0)), bold=True, fill=_status_color(status), color=TRAFFIC_COLORS["white"] if status != "Partially Compliant" else TRAFFIC_COLORS["black"])


def _add_kpi_summary_table(document: Document, kpis: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=5)
    for idx, header in enumerate(["KPI", "Current", "Reason", "Target", "Status"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    if kpis.empty:
        cells = table.add_row().cells
        _set_cell(cells[0], "No KPI data available")
        _set_cell(cells[1], "Not Established", fill=TRAFFIC_COLORS["grey"])
        _set_cell(cells[2], "No KPI framework or KPI-specific evidence was identified in the uploaded documents.")
        _set_cell(cells[3], "-")
        _set_cell(cells[4], "Not Established", fill=TRAFFIC_COLORS["grey"])
    for _, row_data in kpis.iterrows():
        cells = table.add_row().cells
        current = row_data.get("Current")
        target = row_data.get("Target")
        unit = str(row_data.get("Unit", "%"))
        status = str(row_data.get("Status", "Not Established"))
        _set_cell(cells[0], row_data.get("KPI", ""))
        if current is None or pd.isna(current):
            _set_cell(cells[1], row_data.get("Current Display", "Not Established"), fill=TRAFFIC_COLORS["grey"])
        else:
            _set_cell(cells[1], _format_kpi_value(float(current), unit), fill=_score_color(int(current)) if unit == "%" else TRAFFIC_COLORS["light_grey"])
        _set_cell(cells[2], row_data.get("Current Value Reason", "Current value status was determined from uploaded evidence."))
        _set_cell(cells[3], row_data.get("Target Display", _format_kpi_value(target, unit)), fill=TRAFFIC_COLORS["light_grey"])
        _set_cell(cells[4], status, fill=TRAFFIC_COLORS["grey"] if status != "Measured" else TRAFFIC_COLORS["light_green"])
    _format_summary_table(table)


def _add_roadmap_summary(document: Document, findings: list[dict]) -> None:
    phase_map = {
        "Phase 1": {"timelines": {"1 Month", "30 Days"}, "items": []},
        "Phase 2": {"timelines": {"3 Months", "60 Days"}, "items": []},
        "Phase 3": {"timelines": {"6 Months", "12 Months", "90 Days", "180 Days"}, "items": []},
    }
    for item in findings:
        timeline = str(item.get("Timeline", ""))
        action = _limit_words(str(item.get("Recommendation", "")), 14)
        priority = str(item.get("Priority", "Medium"))
        for phase, config in phase_map.items():
            if timeline in config["timelines"]:
                config["items"].append((priority, action, timeline))
                break
    table = document.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Phase", "Timeline", "Priority", "Action"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for phase, config in phase_map.items():
        default_timeline = {"Phase 1": "30 Days", "Phase 2": "60 Days", "Phase 3": "90 Days"}[phase]
        items = config["items"][:3] or [("Medium", "No action assigned.", default_timeline)]
        for priority, action, timeline in items:
            cells = table.add_row().cells
            _set_cell(cells[0], phase)
            _set_cell(cells[1], _timeline_label(timeline))
            _set_cell(cells[2], priority, bold=True, fill=_roadmap_color(priority))
            _set_cell(cells[3], action)
    _format_summary_table(table)


def _timeline_label(timeline: str) -> str:
    return {
        "1 Month": "30 Days",
        "3 Months": "60 Days",
        "6 Months": "90 Days",
        "12 Months": "90 Days",
        "180 Days": "90 Days",
    }.get(str(timeline), str(timeline))


def _load_font(size: int, bold: bool = False):
    try:
        from PIL import ImageFont

        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                continue
        return ImageFont.load_default()
    except Exception:
        return None


def _draw_bar_chart(path, title: str, data: list[tuple]) -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    width = 1100
    row_h = 50
    left = 310
    right = 120
    top = 72
    height = top + row_h * max(1, len(data)) + 36
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _load_font(24, True) or ImageFont.load_default()
    label_font = _load_font(17, False) or ImageFont.load_default()
    small_font = _load_font(15, False) or ImageFont.load_default()
    draw.text((18, 18), title, fill="#1F4E78", font=title_font)
    max_bar = width - left - right
    for idx, (label, current, target) in enumerate(data):
        y = top + idx * row_h
        draw.text((18, y + 10), label[:34], fill="#000000", font=label_font)
        draw.rectangle((left, y + 9, left + max_bar, y + 25), outline="#D9D9D9", fill="#F2F2F2")
        draw.rectangle((left, y + 31, left + max_bar, y + 44), outline="#D9D9D9", fill="#F2F2F2")
        draw.rectangle((left, y + 9, left + int(max_bar * current / 100), y + 25), fill=f"#{_score_color(current)}")
        draw.rectangle((left, y + 31, left + int(max_bar * target / 100), y + 44), fill="#5B9BD5")
        draw.text((left + max_bar + 12, y + 5), f"{current}%", fill="#000000", font=small_font)
        draw.text((left + max_bar + 12, y + 26), f"{target}%", fill="#000000", font=small_font)
    image.save(path)
    return True


def _add_chart_or_fallback(document: Document, title: str, data: list[tuple[str, int, int]], filename: str) -> None:
    import tempfile
    from pathlib import Path

    if not data and "KPI" in title:
        document.add_paragraph("KPI values not measured from uploaded evidence.")
        return

    path = Path(tempfile.gettempdir()) / filename
    if _draw_bar_chart(path, title, data):
        document.add_picture(str(path), width=Inches(6.7))
    else:
        fallback = pd.DataFrame(data, columns=["Metric", "Current", "Target"])
        _add_table(document, fallback, column_widths=[3.0, 1.0, 1.0])


def _set_multiline_cell(cell, lines: list[str], fill: str | None = None, color: str | None = None) -> None:
    cell.text = ""
    if fill:
        _shade_cell(cell, fill)
    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(str(line))
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _add_top_list_table(document: Document, findings: list[dict]) -> None:
    _add_top_findings_table(document, findings)
    _add_top_recommendations_table(document, findings)


def _add_top_findings_table(document: Document, findings: list[dict]) -> None:
    _add_summary_heading(document, "Top 5 Findings", 1)
    table = document.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Finding / Evidence / Risk Impact", "Score", "Risk"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for item in findings[:5]:
        cells = table.add_row().cells
        risk = str(item.get("Risk", item.get("Priority", "Medium")))
        evidence = "; ".join(item.get("Missing Evidence", [])[:3]) or "No sufficient evidence identified"
        _set_multiline_cell(
            cells[0],
            [
                f"Finding: {_clean_line(item.get('Finding', ''), 14)}",
                f"Evidence: {_clean_line(evidence, 14)}",
                f"Risk Impact: {_clean_line(item.get('Business Impact', 'Control gap may reduce assurance.'), 15)}",
            ],
        )
        _set_cell(cells[1], item.get("Score", "0/5"), bold=True)
        _set_cell(cells[2], risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
    _format_summary_table(table)


def _add_top_recommendations_table(document: Document, findings: list[dict]) -> None:
    _add_summary_heading(document, "Top 5 Recommendations", 1)
    table = document.add_table(rows=1, cols=5)
    for idx, header in enumerate(["Issue", "Recommended Action", "Expected Evidence", "Priority", "Timeline"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for item in findings[:5]:
        cells = table.add_row().cells
        priority = str(item.get("Priority", "Medium"))
        actions = item.get("Recommendation Bullets", [])[:1]
        if not actions:
            actions = [_criterion_action(str(item.get("Title", ""))) or "Approve the required control, assign an accountable owner, and retain evidence of recurring management review."]
        _set_cell(cells[0], _clean_line(item.get("Title", item.get("Finding", "Assessment issue")), 10))
        _set_multiline_cell(cells[1], [_complete_sentence(action, 24) for action in actions if _complete_sentence(action, 24)])
        evidence = "; ".join(item.get("Missing Evidence", [])[:3]) or "Approved control evidence"
        _set_cell(cells[2], _clean_line(evidence, 14))
        _set_cell(cells[3], priority, bold=True, fill=_roadmap_color(priority))
        _set_cell(cells[4], item.get("Timeline", "90 Days"))
    _format_summary_table(table)


def _assessment_basis_label(value: object) -> str:
    text = str(value or "").strip().lower()
    if any(term in text for term in ["ai", "openai", "llm"]):
        return "AI-Assisted Evidence Review"
    return "Document and Evidence Review"


def _add_scorecard_table(document: Document, assessments: dict) -> None:
    table = document.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Module", "Score", "Risk", "Assessment Basis"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for module_name, data in _all_module_rows(assessments):
        score = int(data.get("score", 0) or 0)
        risk = str(data.get("risk_level", data.get("status", "Medium")))
        cells = table.add_row().cells
        _set_cell(cells[0], module_name)
        _set_cell(cells[1], f"{score}/100", bold=True, fill=_score_color(score))
        _set_cell(cells[2], risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
        _set_cell(cells[3], _assessment_basis_label(data.get("assessment_mode")))
    _format_summary_table(table)


def _add_risk_heat_table(document: Document, findings: list[dict]) -> None:
    counts = {risk: 0 for risk in ["Critical", "High", "Medium", "Low"]}
    for item in findings:
        priority = _normalize_priority(item.get("Priority", "Medium"))
        counts[priority] = counts.get(priority, 0) + 1
    total = sum(counts.values()) or 1
    table = document.add_table(rows=1, cols=4)
    for idx, risk in enumerate(["Critical", "High", "Medium", "Low"]):
        _set_cell(table.cell(0, idx), risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
    row = table.add_row().cells
    for idx, risk in enumerate(["Critical", "High", "Medium", "Low"]):
        count = counts.get(risk, 0)
        _set_cell(row[idx], f"{count} ({round((count / total) * 100)}%)", bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
    _format_summary_table(table)


def _add_compliance_summary_table(document: Document, assessments: dict) -> None:
    counts = {"Compliant": 0, "Partially Compliant": 0, "Non-Compliant": 0}
    for _, data in _all_module_rows(assessments):
        matrix = data.get("evidence_matrix")
        if isinstance(matrix, pd.DataFrame) and not matrix.empty and "Compliance Status" in matrix:
            for status, count in matrix["Compliance Status"].value_counts().to_dict().items():
                counts[str(status)] = counts.get(str(status), 0) + int(count)
    table = document.add_table(rows=1, cols=3)
    for idx, status in enumerate(["Compliant", "Partially Compliant", "Non-Compliant"]):
        _set_cell(table.cell(0, idx), status, bold=True, fill=_status_color(status), color=TRAFFIC_COLORS["white"] if status != "Partially Compliant" else TRAFFIC_COLORS["black"])
    row = table.add_row().cells
    for idx, status in enumerate(["Compliant", "Partially Compliant", "Non-Compliant"]):
        count = counts.get(status, 0)
        fill = TRAFFIC_COLORS["grey"] if count == 0 else _status_color(status)
        _set_cell(row[idx], str(count), bold=True, fill=fill, color=TRAFFIC_COLORS["white"] if status != "Partially Compliant" and count else TRAFFIC_COLORS["black"])
    _format_summary_table(table)


def _add_kpi_summary_table(document: Document, kpis: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=7)
    for idx, header in enumerate(["KPI", "Current Value", "Reason", "Target", "Maturity", "Risk", "Status"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    if kpis.empty:
        cells = table.add_row().cells
        _set_cell(cells[0], "KPI Framework")
        _set_cell(cells[1], "Not Established", fill=TRAFFIC_COLORS["grey"])
        _set_cell(cells[2], "No KPI framework or KPI-specific evidence was identified in the uploaded documents.")
        _set_cell(cells[3], "Not Established")
        _set_cell(cells[4], "0/5 - Not Established")
        _set_cell(cells[5], "High", fill=_risk_color("High"), color=TRAFFIC_COLORS["white"])
        _set_cell(cells[6], "Not Established", fill=TRAFFIC_COLORS["grey"])
    for _, row_data in kpis.iterrows():
        cells = table.add_row().cells
        status = str(row_data.get("Status", "Not Established"))
        _set_cell(cells[0], row_data.get("KPI", ""))
        current_display = str(row_data.get("Current Display", "Not Established"))
        _set_cell(cells[1], current_display, fill=TRAFFIC_COLORS["light_grey"])
        _set_cell(cells[2], row_data.get("Current Value Reason", "Current value status was determined from uploaded evidence."))
        _set_cell(cells[3], row_data.get("Target Display", "Not Established"), fill=TRAFFIC_COLORS["light_grey"])
        maturity_score = int(row_data.get("Maturity Score", 0) or 0)
        _set_cell(cells[4], f"{maturity_score}/5 - {format_kpi_maturity(maturity_score).split(' - ', 1)[1]}")
        risk = str(row_data.get("Risk", "High"))
        _set_cell(cells[5], risk, bold=True, fill=_risk_color(risk), color=TRAFFIC_COLORS["white"] if risk in {"Critical", "High"} else TRAFFIC_COLORS["black"])
        _set_cell(cells[6], status, fill=TRAFFIC_COLORS["grey"] if status != "Measured" else TRAFFIC_COLORS["light_green"])
    _format_summary_table(table)


def _add_roadmap_summary(document: Document, findings: list[dict]) -> None:
    phase_map = {
        "Phase 1": {"priority": "Critical", "timeline": "30 Days", "items": []},
        "Phase 2": {"priority": "High", "timeline": "60 Days", "items": []},
        "Phase 3": {"priority": "Medium", "timeline": "90 Days", "items": []},
    }
    for item in findings:
        priority = _normalize_priority(item.get("Priority", "Medium"))
        if priority == "Low":
            priority = "Medium"
        phase = {"Critical": "Phase 1", "High": "Phase 2", "Medium": "Phase 3"}[priority]
        actions = item.get("Recommendation Bullets", [])
        action = _complete_sentence(actions[0] if actions else _criterion_action(str(item.get("Title", ""))), 24)
        if action:
            phase_map[phase]["items"].append(action)
    table = document.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Phase", "Timeline", "Priority", "Action"]):
        _set_cell(table.cell(0, idx), header, bold=True, color=TRAFFIC_COLORS["white"])
    for phase, config in phase_map.items():
        actions = list(dict.fromkeys(config["items"]))[:3]
        if not actions:
            actions = [f"No {config['priority'].lower()}-priority findings require action in this phase."]
        cells = table.add_row().cells
        _set_cell(cells[0], phase)
        _set_cell(cells[1], config["timeline"])
        _set_cell(cells[2], config["priority"], bold=True, fill=_roadmap_color(config["priority"]))
        _set_cell(cells[3], " ".join(actions))
    _format_summary_table(table)


def _draw_bar_chart(path, title: str, data: list[tuple[str, int, int]]) -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False
    width = 1100
    row_h = 50
    left = 310
    right = 120
    top = 72
    height = top + row_h * max(1, len(data)) + 46
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _load_font(24, True) or ImageFont.load_default()
    label_font = _load_font(17, False) or ImageFont.load_default()
    small_font = _load_font(15, False) or ImageFont.load_default()
    draw.text((18, 18), title, fill="#1F4E78", font=title_font)
    max_bar = width - left - right
    if not data:
        message = "KPI values not measured from uploaded evidence." if "KPI" in title else "No measured data available."
        draw.text((18, 86), message, fill="#666666", font=label_font)
        image.save(path)
        return True
    values = []
    for item in data:
        values.extend([float(item[1]), float(item[2])])
    scale_max = 100 if all(len(item) < 4 or item[3] == "%" for item in data) else max(values) * 1.25
    scale_max = max(scale_max, 1)
    for idx, item in enumerate(data):
        label, current, target = str(item[0]), float(item[1]), float(item[2])
        unit = item[3] if len(item) > 3 else "%"
        y = top + idx * row_h
        draw.text((18, y + 10), label[:34], fill="#000000", font=label_font)
        draw.rectangle((left, y + 9, left + max_bar, y + 25), outline="#D9D9D9", fill="#F2F2F2")
        draw.rectangle((left, y + 31, left + max_bar, y + 44), outline="#D9D9D9", fill="#F2F2F2")
        draw.rectangle((left, y + 9, left + int(max_bar * current / scale_max), y + 25), fill=f"#{_score_color(int(current)) if unit == '%' else TRAFFIC_COLORS['amber']}")
        draw.rectangle((left, y + 31, left + int(max_bar * target / scale_max), y + 44), fill="#5B9BD5")
        draw.text((left + max_bar + 12, y + 5), _format_kpi_value(current, unit), fill="#000000", font=small_font)
        draw.text((left + max_bar + 12, y + 26), _format_kpi_value(target, unit), fill="#000000", font=small_font)
    image.save(path)
    return True


def build_client_executive_summary_docx(client_profile: dict, documents: list[dict], assessments: dict) -> BytesIO:
    document = Document()
    _configure_document(document)
    for section in document.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    document.styles["Normal"].font.size = Pt(8.5)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Combined Executive Summary")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(TRAFFIC_COLORS["navy"])

    client_name = client_profile.get("client_name") or "Client"
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run(f"{client_name} | Integrated Capability Assessment")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(9)

    overall = _overall_score(assessments)
    risk = _highest_risk(assessments)
    compliance = _compliance_percent(assessments)
    findings = _summary_findings(assessments)
    critical_count = sum(1 for item in findings if item.get("Priority") == "Critical")

    _add_card_row(
        document,
        [
            ("Overall Score", f"{overall}%", _score_color(overall)),
            ("Risk Level", risk, _risk_color(risk)),
            ("Compliance", f"{compliance}%", _score_color(compliance)),
            ("Critical Findings", str(critical_count), _risk_color("Critical")),
        ],
    )

    _add_summary_heading(document, "Executive Summary", 1)
    document.add_paragraph(_portfolio_business_summary(assessments, findings))
    if findings:
        issue_paragraph = document.add_paragraph()
        issue_run = issue_paragraph.add_run("Top Critical Issues: ")
        issue_run.bold = True
        issue_paragraph.add_run("; ".join(str(item.get("Title", "Assessment issue")) for item in findings[:5]))

    _add_summary_heading(document, "Executive Scorecard", 1)
    _add_scorecard_table(document, assessments)
    _add_summary_heading(document, "Risk Heat Map", 1)
    _add_risk_heat_table(document, findings)
    _add_summary_heading(document, "Compliance Dashboard", 1)
    _add_compliance_summary_table(document, assessments)

    _add_top_list_table(document, findings)

    document.add_page_break()
    kpis = _summary_kpis(assessments)
    _add_summary_heading(document, "KPI Dashboard", 1)
    if kpis.attrs.get("placeholder_warning"):
        document.add_paragraph(kpis.attrs["placeholder_warning"])
    _add_kpi_summary_table(document, kpis)
    _add_summary_heading(document, "KPI Charts", 1)
    kpi_chart_data = [
        (str(row["KPI"]), float(row["Current"]), float(row["Target"]), str(row.get("Unit", "%")))
        for _, row in kpis.head(8).iterrows()
        if row["Current"] is not None
        and not pd.isna(row["Current"])
        and row["Target"] is not None
        and not pd.isna(row["Target"])
    ]
    _add_chart_or_fallback(document, "KPI Current vs Target", kpi_chart_data, "client_summary_kpi_chart.png")
    _add_summary_heading(document, "Maturity Chart", 1)
    _add_chart_or_fallback(document, "Maturity Current vs Target", _maturity_rows(assessments), "client_summary_maturity_chart.png")

    _add_summary_heading(document, "Phase Roadmap", 1)
    _add_roadmap_summary(document, findings)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_combined_executive_summary_docx(client_profile: dict, documents: list[dict], assessments: dict) -> BytesIO:
    return build_client_executive_summary_docx(client_profile, documents, assessments)
